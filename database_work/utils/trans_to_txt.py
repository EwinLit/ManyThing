#pip install PyMuPDF
#pip install python-docx
#pip install Spire.Doc
import fitz
import os
from docx import Document as DocxDocument
from spire.doc import Document, FileFormat
from pathlib import Path  # 确保导入 Path
def pdf2txt(pdf_file):
    """Convert all .pdf files in the specified folder and its subfolders to .txt files."""
    
    if not os.path.exists(pdf_file):
        print("指定的文件不存在。")
        return
    
    
    pdf_document = fitz.open(pdf_file)
        
        # 获取 PDF 中的所有页面并将它们合并为一个字符串
    pdf_text = ""
    for page_num in range(pdf_document.page_count):
        page = pdf_document.load_page(page_num)
        pdf_text += page.get_text()
    
    # 创建输出文件路径（同一目录下的 .txt 文件）
    txt_path = os.path.splitext(pdf_file)[0] + '.txt'
    
    # 创建一个 TXT 文件并将 PDF 内容写入其中
    with open(txt_path, 'w', encoding='utf-8') as txt_file:
        txt_file.write(pdf_text)
    
    pdf_document.close()
    
    print(f"成功转换 {pdf_file} 到 {txt_path}")

    #print("所有 PDF 文件已成功转换为 UTF-8 编码的 TXT 文件。")

def convert_docx_to_txt(docx_file):
    """
    将指定的 .docx 文件转换为 TXT 文件，并保存在同一目录下。

    :param docx_file: 指定的 .docx 文件路径。
    """
    if not os.path.exists(docx_file):
        print("指定的文件不存在。")
        return
    
    try:
        # 构造对应的 .txt 文件路径
        txt_path = os.path.splitext(docx_file)[0] + '.txt'
        doc=DocxDocument(docx_file)
        # 提取所有段落的文本
        docx_text = "\n".join([para.text for para in doc.paragraphs])
        
        # 创建一个 TXT 文件并将 DOCX 内容写入其中
        with open(txt_path, 'w', encoding='utf-8') as txt_file:
            txt_file.write(docx_text)

        # 使用 docx2txt 提取文本
        #docx_text = docx_process(docx_file)
        
        # 创建一个 TXT 文件并将 DOCX 内容写入其中
        
        
        print(f"成功转换 {docx_file} 到 {txt_path}")
    
    except Exception as e:
        print(f"转换 {docx_file} 时出错: {e}")
def remove_evaluation_warning(text):
    '''移除水印'''
    lines=text.splitlines()
    if "Evaluation Warning: "in lines[0]:
        return "\n".join(lines[1:])
    return text
def doc2txt(doc_file):
    '''包含水印'''
    
    #print("start transfering")
    txt_path=os.path.splitext(doc_file)[0]+'.txt'
    doc=Document(doc_file)
    doc.LoadFromFile(doc_file)
    doc.SaveToFile(txt_path,FileFormat.Txt)
    doc.Close()
    with open(txt_path,'r',encoding='utf-8')as f:
        content=f.read()
    clean_content=remove_evaluation_warning(content)
    with open(txt_path,'w',encoding='utf-8')as f:
        f.write(clean_content)
   

def convert_file_to_txt(input_file):
    """
    根据文件后缀名选择调用相应的转换函数。

    :param input_file: 输入文件路径。
    """
    file_extension = Path(input_file).suffix.lower()
    
    if file_extension == '.pdf':
        pdf2txt(input_file)
    elif file_extension == '.docx':
        convert_docx_to_txt(input_file)
    elif file_extension == '.doc':
        doc2txt(input_file)
    else:
        print(f"不支持的文件类型: {file_extension}")
if __name__=="__main__":
    INPUTDIR="/home/ubuntu/jiah/workspace/ManyThing/database_work/test/sad.doc"
    convert_file_to_txt(INPUTDIR)
